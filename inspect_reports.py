#!/usr/bin/env python3
import sys
KEYS = ["Shailendra", "Confidential", "Prepared by", "Prepared for", "Generated for", "internal engineering"]
PDFS = [
    r"C:\Users\asus\Desktop\Netweb_100G_RoCE_Benchmark_Report.pdf",
    r"C:\Users\asus\Desktop\Netweb_217_3h_Memory_Validation_Report.pdf",
    r"C:\Users\asus\Desktop\Netweb_Memory_Validation_Report.pdf",
    r"C:\Users\asus\Desktop\Netweb_Server19_Memory_Validation_Report.pdf",
]
DOCX = r"C:\Users\asus\Desktop\Netweb_25G_NIC_Bonding_Benchmark_Report.docx"

import fitz
for p in PDFS:
    print("\n==== PDF:", p.split("\\")[-1], "====")
    d = fitz.open(p)
    print("pages:", d.page_count)
    for pno in range(d.page_count):
        for ln in d[pno].get_text("text").splitlines():
            if any(kk.lower() in ln.lower() for kk in KEYS):
                print(f"  [p{pno+1}] {ln.strip()[:160]}")
    d.close()

import docx
print("\n==== DOCX:", DOCX.split("\\")[-1], "====")
doc = docx.Document(DOCX)
for i, para in enumerate(doc.paragraphs):
    if any(kk.lower() in para.text.lower() for kk in KEYS):
        print(f"  [body para {i}] {para.text.strip()[:180]}")
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cl in enumerate(row.cells):
            if any(kk.lower() in cl.text.lower() for kk in KEYS):
                print(f"  [tbl{ti} r{ri} c{ci}] {cl.text.strip()[:180]}")
for si, sec in enumerate(doc.sections):
    for kind, hf in (("footer", sec.footer), ("header", sec.header), ("first_footer", sec.first_page_footer), ("first_header", sec.first_page_header)):
        try:
            for para in hf.paragraphs:
                if any(kk.lower() in para.text.lower() for kk in KEYS):
                    print(f"  [sec{si} {kind}] {para.text.strip()[:180]}")
        except Exception:
            pass
