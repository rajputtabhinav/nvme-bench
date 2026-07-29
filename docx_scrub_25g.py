#!/usr/bin/env python3
import docx, shutil, os
P = r"C:\Users\asus\Desktop\Netweb_25G_NIC_Bonding_Benchmark_Report.docx"
BK = r"C:\Users\asus\Desktop\nvme_bench\report_backups"
os.makedirs(BK, exist_ok=True)
shutil.copy2(P, os.path.join(BK, os.path.basename(P)))
doc = docx.Document(P)


def repl(par, old, new):
    if old not in par.text:
        return 0
    cnt = 0
    handled = False
    for r in par.runs:
        if old in r.text:
            r.text = r.text.replace(old, new); cnt += 1; handled = True
    if not handled:  # spans runs -> collapse into first run
        joined = "".join(r.text for r in par.runs).replace(old, new)
        for r in par.runs:
            r.text = ""
        if par.runs:
            par.runs[0].text = joined
        cnt += 1
    return cnt


def scrub_par(par):
    n = 0
    n += repl(par, "Shailendra Rajput", "Netweb Technologies India Ltd")
    n += repl(par, "   |   Confidential", "")
    n += repl(par, "| Confidential", "")
    n += repl(par, "Confidential", "")
    return n

total = 0
for par in doc.paragraphs:
    total += scrub_par(par)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                total += scrub_par(par)
for sec in doc.sections:
    for hf in (sec.first_page_footer, sec.footer, sec.even_page_footer,
               sec.first_page_header, sec.header, sec.even_page_header):
        try:
            for par in hf.paragraphs:
                total += scrub_par(par)
        except Exception:
            pass
doc.save(P)
print("DOCX_DONE edits:", total)
